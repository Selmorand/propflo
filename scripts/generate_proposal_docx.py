"""
Regenerate PropFlo-Full-Document-Package.docx from the source markdown files.

Sources (in order):
    02_Proposal/Proposal-Draft.md
    02_Proposal/Scope-of-Work.md
    02_Proposal/Implementation-Plan.md
    02_Proposal/Assumptions-and-Limitations.md

Behaviour:
    * Cover page with title, subtitle, and prepared-for/by line.
    * Contents page listing the four documents.
    * Each source rendered under a Heading 1 "Part N. <Document title>".
    * Within Proposal-Draft.md, sections 1-3 (Proposal Title, Prepared For, Prepared By)
      are dropped because that metadata sits on the cover page. Remaining sections are
      renumbered so the visible sequence starts at 1.
    * Other three files keep their original numbering.
    * Handles headings (## / ###), tables (| ... |), bullets, bold (**...**), and
      backticks (rendered as inline code style).

Run:
    python scripts/generate_proposal_docx.py

Requires:
    python-docx (pip install python-docx)
"""
from __future__ import annotations

import os
import re
import sys
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# -------------------------------------------------------------- config

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

SOURCES = [
    ("Marketing Proposal", os.path.join("02_Proposal", "Proposal-Draft.md"), True),
    ("Scope of Work", os.path.join("02_Proposal", "Scope-of-Work.md"), False),
    ("Implementation Plan", os.path.join("02_Proposal", "Implementation-Plan.md"), False),
    ("Assumptions and Limitations", os.path.join("02_Proposal", "Assumptions-and-Limitations.md"), False),
]

OUTPUT = os.path.join(REPO_ROOT, "PropFlo-Full-Document-Package.docx")

# Metadata sections at the top of Proposal-Draft.md that live on the cover page instead.
DROP_PROPOSAL_SECTIONS = {"1. Proposal Title", "2. Prepared For", "3. Prepared By"}


# -------------------------------------------------------------- lightweight markdown AST

@dataclass
class Block:
    kind: str  # "h1" "h2" "h3" "p" "bullet" "table" "blank"
    text: str = ""
    rows: List[List[str]] = field(default_factory=list)  # only for tables


def parse_markdown(text: str) -> List[Block]:
    lines = text.splitlines()
    blocks: List[Block] = []
    i = 0

    def flush_paragraph(buf: List[str]) -> None:
        if buf:
            blocks.append(Block("p", " ".join(s.strip() for s in buf)))
            buf.clear()

    para_buf: List[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if stripped.startswith("| ") and "|" in stripped[2:]:
            # table starts
            flush_paragraph(para_buf)
            rows: List[List[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # skip alignment separator row like |---|---|
                if all(re.match(r"^:?-+:?$", c) for c in row if c):
                    i += 1
                    continue
                rows.append(row)
                i += 1
            if rows:
                blocks.append(Block("table", rows=rows))
            continue

        if stripped.startswith("- "):
            flush_paragraph(para_buf)
            blocks.append(Block("bullet", stripped[2:].strip()))
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph(para_buf)
            blocks.append(Block("h3", stripped[4:].strip()))
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph(para_buf)
            blocks.append(Block("h2", stripped[3:].strip()))
            i += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph(para_buf)
            blocks.append(Block("h1", stripped[2:].strip()))
            i += 1
            continue

        if stripped == "":
            flush_paragraph(para_buf)
            blocks.append(Block("blank"))
            i += 1
            continue

        para_buf.append(stripped)
        i += 1

    flush_paragraph(para_buf)
    return blocks


# -------------------------------------------------------------- inline formatting

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_runs(paragraph, text: str) -> None:
    """Add text to paragraph, honouring **bold** inline spans. Backticks are stripped
    to match the plain rendering used in the previous docx."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            paragraph.add_run(part[1:-1])
        else:
            paragraph.add_run(part)


# -------------------------------------------------------------- table helper

def render_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Normal Table"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = tbl.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


# -------------------------------------------------------------- proposal section renumbering

SECTION_HEADING_RE = re.compile(r"^(\d+)\.\s+(.+)$")


def renumber_proposal_blocks(blocks: List[Block]) -> List[Block]:
    """Drop metadata sections 1-3, then renumber remaining ## X. headings sequentially."""
    output: List[Block] = []
    skip_until_next_h2 = False
    next_number = 1

    for block in blocks:
        if block.kind == "h2":
            m = SECTION_HEADING_RE.match(block.text)
            if m and block.text in DROP_PROPOSAL_SECTIONS:
                skip_until_next_h2 = True
                continue
            skip_until_next_h2 = False
            if m:
                _, title = m.groups()
                block = Block("h2", f"{next_number}. {title}")
                next_number += 1
            output.append(block)
            continue

        if block.kind in ("h1",):
            # H1 in Proposal-Draft.md is the document title; skip (rendered as Part heading)
            skip_until_next_h2 = False
            continue

        if skip_until_next_h2:
            continue

        output.append(block)

    return output


def strip_document_title(blocks: List[Block]) -> List[Block]:
    """Drop the single top-level H1 heading; keep everything else."""
    return [b for b in blocks if b.kind != "h1"]


# -------------------------------------------------------------- top-level rendering

def render_cover(doc: Document) -> None:
    """Cover page uses Normal paragraphs with sized runs, matching the previous docx."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PropFlo Marketing Proposal")
    r.bold = True
    r.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Complete Document Package")
    r.font.size = Pt(16)

    doc.add_paragraph()

    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = prep.add_run("Prepared for PropFlo  |  Prepared by Interon")
    r.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    contents = doc.add_paragraph()
    contents.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = contents.add_run("Contents")
    r.bold = True
    r.font.size = Pt(14)

    for i, (title_text, _, _) in enumerate(SOURCES, 1):
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.add_run(f"{i}. {title_text}").font.size = Pt(12)

    doc.add_page_break()


def render_blocks(doc: Document, blocks: List[Block]) -> None:
    for block in blocks:
        if block.kind == "blank":
            continue
        if block.kind == "h2":
            para = doc.add_paragraph(style="Heading 1")
            add_runs(para, block.text)
        elif block.kind == "h3":
            para = doc.add_paragraph(style="Heading 2")
            add_runs(para, block.text)
        elif block.kind == "bullet":
            para = doc.add_paragraph(style="List Paragraph")
            add_runs(para, block.text)
        elif block.kind == "p":
            para = doc.add_paragraph()
            add_runs(para, block.text)
        elif block.kind == "table":
            render_table(doc, block.rows)
            doc.add_paragraph()


def render_source(doc: Document, part_index: int, title: str, path: str, is_proposal: bool) -> None:
    """Part heading rendered as a plain centred Normal paragraph, matching the previous docx."""
    part_heading = doc.add_paragraph()
    part_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = part_heading.add_run(f"Part {part_index}: {title}")
    r.bold = True
    r.font.size = Pt(20)
    doc.add_paragraph()

    full_path = os.path.join(REPO_ROOT, path)
    with open(full_path, encoding="utf-8") as f:
        blocks = parse_markdown(f.read())

    if is_proposal:
        blocks = renumber_proposal_blocks(blocks)
    else:
        blocks = strip_document_title(blocks)

    render_blocks(doc, blocks)

    doc.add_page_break()


# -------------------------------------------------------------- entry point

def main() -> None:
    doc = Document()

    render_cover(doc)
    for i, (title, path, is_proposal) in enumerate(SOURCES, 1):
        render_source(doc, i, title, path, is_proposal)

    try:
        doc.save(OUTPUT)
        print(f"Wrote {OUTPUT} ({os.path.getsize(OUTPUT):,} bytes)")
    except PermissionError:
        staging = OUTPUT + ".new"
        doc.save(staging)
        print(
            f"Target file is locked (probably open in Word). Wrote staging copy:\n"
            f"  {staging} ({os.path.getsize(staging):,} bytes)\n"
            f"Close Word, then rename to replace the original."
        )


if __name__ == "__main__":
    main()
